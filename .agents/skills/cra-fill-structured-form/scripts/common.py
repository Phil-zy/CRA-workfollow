from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from uuid import uuid4
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook
from jsonschema import Draft202012Validator


ALLOWED_STATUSES = {"已填写", "已核对", "缺失", "冲突", "待确认", "人工保留"}
MISSING_SENTINELS = {"", "未勾选"}
CHECKLIST_HEADERS = ["字段名称", "目标位置", "处理状态", "来源值", "目标原有值", "未填写或冲突原因", "来源文件及位置", "建议 CRA 操作", "CRA 最终决定"]
CONFIG_SCHEMA_PATH = Path(__file__).resolve().parents[1] / "assets" / "template-config.schema.json"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def resolved(path: str | Path) -> Path:
    return Path(path).expanduser().resolve(strict=False)


def ensure_within(workspace: Path, paths: Iterable[Path]) -> None:
    workspace = workspace.resolve(strict=True)
    for path in paths:
        candidate = path.resolve(strict=False)
        try:
            candidate.relative_to(workspace)
        except ValueError as exc:
            raise ValueError(f"路径不在授权工作区内: {candidate}") from exc


def ensure_output_isolated(output_dir: Path, input_paths: Iterable[Path]) -> None:
    output_dir = output_dir.resolve(strict=False)
    for input_path in input_paths:
        resolved_input = input_path.resolve(strict=True)
        if output_dir == resolved_input.parent:
            raise ValueError(f"输出目录必须与输入目录隔离: {output_dir}")
        try:
            resolved_input.relative_to(output_dir)
        except ValueError:
            continue
        raise ValueError(f"输入文件不得位于输出目录内: {resolved_input}")


def require_file(path: Path, suffix: str) -> None:
    if not path.is_file():
        raise FileNotFoundError(f"文件不存在: {path}")
    if path.suffix.lower() != suffix:
        raise ValueError(f"文件格式必须为 {suffix}: {path}")


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, date):
        return value.isoformat()
    return re.sub(r"\s+", " ", str(value)).strip()


def display_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, date):
        return value.isoformat()
    return str(value)


def load_facts(path: Path) -> dict[str, dict[str, Any]]:
    require_file(path, ".xlsx")
    workbook = load_workbook(path, data_only=True, read_only=True)
    if not workbook.worksheets:
        raise ValueError("项目事实工作簿没有工作表")
    sheet = workbook.worksheets[0]
    facts: dict[str, dict[str, Any]] = {}
    for row in range(1, sheet.max_row + 1):
        key = normalize_text(sheet.cell(row=row, column=1).value)
        if not key or key == "字段":
            continue
        raw = sheet.cell(row=row, column=2).value
        value = display_value(raw)
        missing = normalize_text(value) in MISSING_SENTINELS
        facts[key] = {
            "key": key,
            "value": value,
            "status": "missing" if missing else "confirmed",
            "source_file": path.name,
            "source_sheet": sheet.title,
            "source_location": f"{sheet.title}!B{row}",
        }
    workbook.close()
    return facts


def _unique_cells(table: Any) -> list[dict[str, Any]]:
    cells: list[dict[str, Any]] = []
    seen: set[Any] = set()
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            identity = cell._tc
            if identity in seen:
                continue
            seen.add(identity)
            grid_span = cell._tc.tcPr.gridSpan
            v_merge = cell._tc.tcPr.vMerge
            cells.append(
                {
                    "row": row_index,
                    "column": column_index,
                    "grid_span": int(grid_span.val) if grid_span is not None else 1,
                    "v_merge": str(v_merge.val) if v_merge is not None else None,
                    "text": cell.text,
                }
            )
    return cells


def inspect_docx(path: Path) -> dict[str, Any]:
    require_file(path, ".docx")
    try:
        with zipfile.ZipFile(path) as package:
            members = set(package.namelist())
            if any(name.endswith("vbaProject.bin") for name in members):
                raise ValueError("目标 Word 包含宏，当前原型禁止处理")
            if "word/settings.xml" in members:
                settings = package.read("word/settings.xml")
                if b"documentProtection" in settings:
                    raise ValueError("目标 Word 已启用文档保护，当前原型禁止处理")
    except zipfile.BadZipFile as exc:
        raise ValueError("目标 Word 无法作为安全的 DOCX 包读取，可能已加密或文件损坏") from exc
    try:
        document = Document(path)
    except Exception as exc:
        raise ValueError("目标 Word 无法安全读取，可能已加密、受保护或文件损坏") from exc
    tables = []
    for table_index, table in enumerate(document.tables):
        tables.append(
            {
                "index": table_index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": _unique_cells(table),
            }
        )
    structure = {
        "paragraphs": [paragraph.text for paragraph in document.paragraphs],
        "sections": len(document.sections),
        "tables": tables,
    }
    serialized = json.dumps(structure, ensure_ascii=False, sort_keys=True).encode("utf-8")
    geometry = {
        "paragraph_count": len(document.paragraphs),
        "sections": len(document.sections),
        "tables": [
            {
                "rows": table["rows"],
                "columns": table["columns"],
                "cells": [
                    {key: cell[key] for key in ("row", "column", "grid_span", "v_merge")}
                    for cell in table["cells"]
                ],
            }
            for table in tables
        ],
    }
    return {
        "file_name": path.name,
        "sha256": sha256_file(path),
        "structure_fingerprint": hashlib.sha256(serialized).hexdigest(),
        "format": "docx",
        "paragraph_count": len(document.paragraphs),
        "section_count": len(document.sections),
        "table_count": len(document.tables),
        "geometry": geometry,
        "structure": structure,
    }


def inspect_xlsx(path: Path) -> dict[str, Any]:
    require_file(path, ".xlsx")
    try:
        with zipfile.ZipFile(path) as package:
            members = set(package.namelist())
            if any(name.endswith("vbaProject.bin") for name in members):
                raise ValueError("目标 Excel 包含宏，当前首版禁止处理")
            if any(name.startswith("xl/externalLinks/") for name in members) or "xl/connections.xml" in members:
                raise ValueError("目标 Excel 包含外部数据连接，当前首版禁止处理")
    except zipfile.BadZipFile as exc:
        raise ValueError("目标 Excel 无法作为安全的 XLSX 包读取，可能已加密或文件损坏") from exc
    try:
        workbook = load_workbook(path, data_only=False, read_only=False, keep_links=True)
    except Exception as exc:
        raise ValueError("目标 Excel 无法安全读取，可能已加密、受保护或文件损坏") from exc
    try:
        security = workbook.security
        if security and (security.lockStructure or security.lockWindows or security.lockRevision):
            raise ValueError("目标 Excel 已启用工作簿保护，当前首版禁止处理")
        if any(sheet.protection.sheet for sheet in workbook.worksheets):
            raise ValueError("目标 Excel 已启用工作表保护，当前首版禁止处理")

        sheets: list[dict[str, Any]] = []
        geometry_sheets: list[dict[str, Any]] = []
        for sheet in workbook.worksheets:
            populated_cells = []
            formula_cells = []
            style_cells = []
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        value = display_value(cell.value)
                        populated_cells.append({"coordinate": cell.coordinate, "value": value, "data_type": cell.data_type})
                        if cell.data_type == "f":
                            formula_cells.append({"coordinate": cell.coordinate, "formula": value})
                    if cell.has_style:
                        style_cells.append({"coordinate": cell.coordinate, "style_id": cell.style_id})
            merged_ranges = sorted(str(cell_range) for cell_range in sheet.merged_cells.ranges)
            row_dimensions = {
                str(index): {
                    "height": dimension.height,
                    "hidden": bool(dimension.hidden),
                    "style_id": dimension.style_id,
                    "outline_level": dimension.outlineLevel,
                    "collapsed": bool(dimension.collapsed),
                }
                for index, dimension in sheet.row_dimensions.items()
                if dimension.height is not None or dimension.hidden or dimension.style_id or dimension.outlineLevel or dimension.collapsed
            }
            column_dimensions = {
                key: {
                    "width": dimension.width,
                    "hidden": bool(dimension.hidden),
                    "style_id": dimension.style_id,
                    "outline_level": dimension.outlineLevel,
                    "collapsed": bool(dimension.collapsed),
                    "best_fit": bool(dimension.bestFit),
                }
                for key, dimension in sheet.column_dimensions.items()
                if dimension.width is not None or dimension.hidden or dimension.style_id or dimension.outlineLevel or dimension.collapsed or dimension.bestFit
            }
            conditional_formatting = []
            for conditional_range in sheet.conditional_formatting:
                rules = sheet.conditional_formatting[conditional_range]
                conditional_formatting.append(
                    {
                        "range": str(conditional_range.sqref),
                        "rules": [
                            {
                                "type": rule.type,
                                "priority": rule.priority,
                                "operator": rule.operator,
                                "formula": list(rule.formula or []),
                                "text": rule.text,
                                "stop_if_true": rule.stopIfTrue,
                                "dxf_id": rule.dxfId,
                            }
                            for rule in rules
                        ],
                    }
                )
            page = {
                "orientation": sheet.page_setup.orientation,
                "paper_size": sheet.page_setup.paperSize,
                "scale": sheet.page_setup.scale,
                "fit_to_width": sheet.page_setup.fitToWidth,
                "fit_to_height": sheet.page_setup.fitToHeight,
                "first_page_number": sheet.page_setup.firstPageNumber,
                "use_first_page_number": sheet.page_setup.useFirstPageNumber,
                "black_and_white": sheet.page_setup.blackAndWhite,
                "draft": sheet.page_setup.draft,
                "horizontal_dpi": sheet.page_setup.horizontalDpi,
                "vertical_dpi": sheet.page_setup.verticalDpi,
                "print_area": str(sheet.print_area or ""),
                "print_title_rows": str(sheet.print_title_rows or ""),
                "print_title_cols": str(sheet.print_title_cols or ""),
                "margins": {
                    "left": sheet.page_margins.left,
                    "right": sheet.page_margins.right,
                    "top": sheet.page_margins.top,
                    "bottom": sheet.page_margins.bottom,
                    "header": sheet.page_margins.header,
                    "footer": sheet.page_margins.footer,
                },
                "print_options": {
                    "horizontal_centered": sheet.print_options.horizontalCentered,
                    "vertical_centered": sheet.print_options.verticalCentered,
                    "headings": sheet.print_options.headings,
                    "grid_lines": sheet.print_options.gridLines,
                    "grid_lines_set": sheet.print_options.gridLinesSet,
                },
                "odd_header": str(sheet.oddHeader),
                "odd_footer": str(sheet.oddFooter),
                "even_header": str(sheet.evenHeader),
                "even_footer": str(sheet.evenFooter),
                "first_header": str(sheet.firstHeader),
                "first_footer": str(sheet.firstFooter),
            }
            sheets.append(
                {
                    "name": sheet.title,
                    "state": sheet.sheet_state,
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                    "merged_ranges": merged_ranges,
                    "populated_cells": populated_cells,
                    "style_cells": style_cells,
                    "row_dimensions": row_dimensions,
                    "column_dimensions": column_dimensions,
                    "conditional_formatting": conditional_formatting,
                    "page": page,
                }
            )
            geometry_sheets.append(
                {
                    "name": sheet.title,
                    "state": sheet.sheet_state,
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                    "merged_ranges": merged_ranges,
                    "formula_cells": formula_cells,
                    "style_cells": style_cells,
                    "row_dimensions": row_dimensions,
                    "column_dimensions": column_dimensions,
                    "conditional_formatting": conditional_formatting,
                    "page": page,
                }
            )
        structure = {"sheets": sheets}
        serialized = json.dumps(structure, ensure_ascii=False, sort_keys=True).encode("utf-8")
        return {
            "file_name": path.name,
            "sha256": sha256_file(path),
            "structure_fingerprint": hashlib.sha256(serialized).hexdigest(),
            "format": "xlsx",
            "sheet_count": len(workbook.worksheets),
            "sheet_names": workbook.sheetnames,
            "geometry": {"sheets": geometry_sheets},
            "structure": structure,
        }
    finally:
        workbook.close()


def inspect_template(path: Path) -> dict[str, Any]:
    if path.suffix.lower() == ".docx":
        return inspect_docx(path)
    if path.suffix.lower() == ".xlsx":
        return inspect_xlsx(path)
    raise ValueError(f"目标文件格式必须为 .docx 或 .xlsx: {path}")


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        value = json.load(handle)
    if not isinstance(value, dict):
        raise ValueError(f"JSON 顶层必须是对象: {path}")
    return value


def write_json_new(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = temporary_output_path(path)
    try:
        with temporary.open("x", encoding="utf-8", newline="\n") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)
            handle.write("\n")
        publish_new_file(temporary, path)
    finally:
        if temporary.exists():
            temporary.unlink()


def matching_configs(config_library: Path, fingerprint: str, sha256: str, status: str) -> list[Path]:
    directory = config_library / ({"draft": "drafts", "enabled": "enabled", "disabled": "disabled"}[status])
    if not directory.is_dir():
        return []
    matches: list[Path] = []
    for path in sorted(directory.glob("*.json")):
        try:
            config = load_json(path)
        except (OSError, ValueError, json.JSONDecodeError):
            continue
        template = config.get("template", {})
        if (
            config.get("status") == status
            and template.get("structure_fingerprint") == fingerprint
            and template.get("sha256") == sha256
        ):
            matches.append(path)
    return matches


def next_config_version(config_library: Path, config_id: str) -> int:
    versions: list[int] = []
    for directory_name in ("drafts", "enabled", "disabled"):
        directory = config_library / directory_name
        if not directory.is_dir():
            continue
        for path in directory.glob("*.json"):
            try:
                config = load_json(path)
            except (OSError, ValueError, json.JSONDecodeError):
                continue
            if config.get("config_id") == config_id and isinstance(config.get("version"), int):
                versions.append(config["version"])
    return max(versions, default=0) + 1


def validate_config(config: dict[str, Any], required_status: str | None = None) -> None:
    schema = load_json(CONFIG_SCHEMA_PATH)
    errors = sorted(Draft202012Validator(schema).iter_errors(config), key=lambda item: list(item.absolute_path))
    if errors:
        details = "; ".join(
            f"{'/'.join(str(part) for part in error.absolute_path) or '<root>'}: {error.message}"
            for error in errors[:5]
        )
        raise ValueError(f"模板配置不符合 Schema: {details}")
    if required_status and config["status"] != required_status:
        raise ValueError(f"模板配置状态必须为 {required_status}")
    ids = [field.get("id") for field in config["fields"]]
    if None in ids or len(ids) != len(set(ids)):
        raise ValueError("模板配置字段 ID 缺失或重复")
    template_format = config["template"]["format"]
    for field in config["fields"]:
        target = field["target"]
        mapping = field["mapping"]
        if template_format == "xlsx":
            if set(target) != {"sheet", "cell", "expected_original"}:
                raise ValueError(f"Excel 配置字段目标必须使用工作表和单元格: {field['id']}")
            if mapping["type"] not in {"direct", "manual"}:
                raise ValueError(f"Excel 首版不支持该映射类型: {field['id']}")
            if mapping["type"] == "direct" and mapping.get("mode") != "cell":
                raise ValueError(f"Excel 直接字段必须使用 cell 模式: {field['id']}")
        elif set(target) != {"table", "row", "column", "expected_original"}:
            raise ValueError(f"Word 配置字段目标必须使用表格、行和列: {field['id']}")
    if config["status"] == "enabled":
        if any(field["mapping_review"] != "approved" for field in config["fields"]):
            raise ValueError("已启用配置存在未批准的字段映射")
        if not all(config["approval"].get(key) for key in ("approved_by", "approved_at", "approval_id")):
            raise ValueError("已启用配置缺少完整 CRA 审批记录")


def update_json_atomic(path: Path, payload: dict[str, Any]) -> None:
    temporary = temporary_output_path(path)
    try:
        with temporary.open("x", encoding="utf-8", newline="\n") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)
            handle.write("\n")
        os.replace(temporary, path)
    finally:
        if temporary.exists():
            temporary.unlink()


def target_cell(document: Any, target: dict[str, Any]) -> Any:
    try:
        return document.tables[int(target["table"])].rows[int(target["row"])].cells[int(target["column"])]
    except (IndexError, KeyError, TypeError, ValueError) as exc:
        raise ValueError(f"目标位置无效: {target}") from exc


def set_blank_cell_text(cell: Any, value: str) -> bool:
    if normalize_text(cell.text):
        return False
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_properties = paragraph._p.get_or_add_pPr()
    indentation = paragraph_properties.find(qn("w:ind"))
    if indentation is not None:
        for attribute in ("firstLine", "firstLineChars", "hanging", "hangingChars"):
            indentation.attrib.pop(qn(f"w:{attribute}"), None)
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = value
    else:
        run = paragraph.add_run(value)
    value_length = len(normalize_text(value))
    if value_length >= 48:
        run.font.size = Pt(9)
    else:
        run.font.size = Pt(10.5)
    return True


def compact_trailing_empty_paragraph(document: Any) -> bool:
    if not document.paragraphs:
        return False
    paragraph = document.paragraphs[-1]
    if normalize_text(paragraph.text):
        return False
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph_properties = paragraph._p.get_or_add_pPr()
    run_properties = paragraph_properties.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        paragraph_properties.append(run_properties)
    for tag_name in ("sz", "szCs"):
        element = run_properties.find(qn(f"w:{tag_name}"))
        if element is None:
            element = OxmlElement(f"w:{tag_name}")
            run_properties.append(element)
        element.set(qn("w:val"), "2")
    return True


def replace_choice(cell: Any, option: str, unchecked: str, checked: str) -> bool:
    pattern = re.compile(re.escape(unchecked) + r"(\s*)" + re.escape(option))
    candidates: list[tuple[list[Any], int]] = []
    for paragraph in cell.paragraphs:
        runs = list(paragraph.runs)
        paragraph_text = "".join(run.text for run in runs)
        for match in pattern.finditer(paragraph_text):
            candidates.append((runs, match.start()))

    if len(candidates) != 1 or len(unchecked) != 1:
        return False

    runs, checkbox_index = candidates[0]
    cursor = 0
    for run in runs:
        next_cursor = cursor + len(run.text)
        if cursor <= checkbox_index < next_cursor:
            offset = checkbox_index - cursor
            if run.text[offset] != unchecked:
                return False
            run.text = run.text[:offset] + checked + run.text[offset + 1 :]
            return True
        cursor = next_cursor
    return False


def selected_options(text: str, options: Iterable[str], checked: str) -> list[str]:
    selected = []
    for option in options:
        if re.search(re.escape(checked) + r"\s*" + re.escape(option), text):
            selected.append(option)
    return selected


def unique_output_path(directory: Path, file_name: str) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    candidate = directory / file_name
    if not candidate.exists():
        return candidate
    stem, suffix = candidate.stem, candidate.suffix
    for index in range(1, 10000):
        alternate = directory / f"{stem}_{index}{suffix}"
        if not alternate.exists():
            return alternate
    raise FileExistsError("无法生成不重名的输出文件")


def temporary_output_path(final_path: Path) -> Path:
    return final_path.with_name(f".{final_path.stem}.{uuid4().hex}.tmp{final_path.suffix}")


def publish_new_file(temporary_path: Path, final_path: Path) -> None:
    """Atomically publish a complete file without ever replacing an existing path."""
    final_path.parent.mkdir(parents=True, exist_ok=True)
    os.link(temporary_path, final_path)


def target_location(field: dict[str, Any]) -> str:
    target = field["target"]
    if "sheet" in target:
        return f"{target['sheet']}!{target['cell']}"
    return f"表{target['table'] + 1}/行{target['row'] + 1}/列{target['column'] + 1}"


def checklist_row(item: dict[str, Any]) -> list[str]:
    source = " / ".join(part for part in (item.get("source_file", ""), item.get("source_location", "")) if part)
    return [
        str(item.get("field_name", "")),
        str(item.get("target_location", "")),
        str(item.get("status", "")),
        str(item.get("source_value", "")),
        str(item.get("target_original", "")),
        str(item.get("reason", "")),
        source,
        str(item.get("suggested_action", "")),
        str(item.get("cra_final_decision", "")),
    ]
