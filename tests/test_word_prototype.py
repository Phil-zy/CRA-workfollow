from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / ".agents" / "skills" / "cra-fill-structured-form" / "scripts"
sys.path.insert(0, str(SCRIPTS))
import common as common_module
import fill_docx as fill_docx_module


def run_script(name: str, *args: object, expect_success: bool = True) -> subprocess.CompletedProcess[str]:
    environment = os.environ.copy()
    environment["PYTHONIOENCODING"] = "utf-8"
    process = subprocess.run(
        [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        capture_output=True,
        check=False,
        env=environment,
    )
    if expect_success and process.returncode != 0:
        raise AssertionError(f"{name} failed:\nSTDOUT:\n{process.stdout}\nSTDERR:\n{process.stderr}")
    if not expect_success and process.returncode == 0:
        raise AssertionError(f"{name} unexpectedly succeeded:\n{process.stdout}")
    return process


class WordPrototypeTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.workspace = Path(self.temp.name)
        self.target = self.workspace / "template.docx"
        self.facts = self.workspace / "facts.xlsx"
        self.config_library = self.workspace / "template-configs"
        self.output_dir = self.workspace / "output"
        self.audit_dir = self.workspace / "audit"
        self.audit = self.audit_dir / "audit.json"
        self._write_docx(self.target)
        self._write_facts(self.facts)

    def tearDown(self) -> None:
        self.temp.cleanup()

    @staticmethod
    def _write_docx(path: Path, protected: bool = False) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = ""
        if protected:
            protection = OxmlElement("w:documentProtection")
            protection.set(qn("w:edit"), "readOnly")
            protection.set(qn("w:enforcement"), "1")
            document.settings.element.append(protection)
        document.save(path)

    @staticmethod
    def _write_facts(path: Path) -> None:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "项目信息"
        sheet.append(["字段", "内容"])
        sheet.append(["项目名称", "测试项目"])
        workbook.save(path)

    def _inspection(self, target: Path | None = None) -> dict:
        process = run_script(
            "inspect_template.py",
            "--target",
            target or self.target,
            "--authorized-workspace",
            self.workspace,
            "--config-library",
            self.config_library,
        )
        return json.loads(process.stdout)

    def _write_config(self, *, status: str = "enabled", invalid_mapping: bool = False) -> Path:
        inspection = self._inspection()["template"]
        mapping = {} if invalid_mapping else {"type": "direct", "mode": "cell"}
        config = {
            "schema_version": "1.0",
            "config_id": "test-form",
            "version": 1,
            "status": status,
            "form_type": "test-form",
            "template": inspection,
            "fields": [
                {
                    "id": "project_name",
                    "label": "项目名称",
                    "fact_key": "项目名称",
                    "target": {"table": 0, "row": 0, "column": 0, "expected_original": ""},
                    "mapping": mapping,
                    "mapping_review": "approved" if status == "enabled" else "proposed",
                }
            ],
            "approval": {
                "approved_by": "CRA" if status == "enabled" else None,
                "approved_at": "2026-08-13T00:00:00+00:00" if status == "enabled" else None,
                "approval_id": "TEST-APPROVAL" if status == "enabled" else None,
            },
        }
        directory = self.config_library / ({"enabled": "enabled", "draft": "drafts"}[status])
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / "test-form-v1.json"
        path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    def _fill_and_build(self) -> tuple[Path, Path, Path]:
        config = self._write_config()
        fill = run_script(
            "fill_docx.py",
            "--target",
            self.target,
            "--facts",
            self.facts,
            "--config",
            config,
            "--authorized-workspace",
            self.workspace,
            "--output-dir",
            self.output_dir,
            "--audit-json",
            self.audit,
            "--authorization-id",
            "TEST-AUTHORIZATION",
        )
        output_docx = Path(json.loads(fill.stdout)["output_docx"])
        checklist_result = run_script(
            "build_checklist.py",
            "--audit-json",
            self.audit,
            "--authorized-workspace",
            self.workspace,
            "--output-dir",
            self.output_dir,
        )
        checklist = Path(json.loads(checklist_result.stdout)["checklist"])
        return config, output_docx, checklist

    def test_inspection_rejects_protected_and_macro_enabled_word_files(self) -> None:
        protected = self.workspace / "protected.docx"
        self._write_docx(protected, protected=True)
        protected_result = run_script(
            "inspect_template.py",
            "--target",
            protected,
            "--authorized-workspace",
            self.workspace,
            "--config-library",
            self.config_library,
            expect_success=False,
        )
        self.assertIn("保护", protected_result.stderr)

        macro = self.workspace / "macro.docx"
        self._write_docx(macro)
        with zipfile.ZipFile(macro, "a") as package:
            package.writestr("word/vbaProject.bin", b"test macro marker")
        macro_result = run_script(
            "inspect_template.py",
            "--target",
            macro,
            "--authorized-workspace",
            self.workspace,
            "--config-library",
            self.config_library,
            expect_success=False,
        )
        self.assertIn("宏", macro_result.stderr)

        encrypted = self.workspace / "encrypted.docx"
        encrypted.write_bytes(b"not-a-zip-encrypted-placeholder")
        encrypted_result = run_script(
            "inspect_template.py",
            "--target",
            encrypted,
            "--authorized-workspace",
            self.workspace,
            "--config-library",
            self.config_library,
            expect_success=False,
        )
        self.assertTrue("加密" in encrypted_result.stderr or "损坏" in encrypted_result.stderr)

    def test_audit_failure_does_not_publish_untracked_word_file(self) -> None:
        config = self._write_config()
        arguments = [
            "fill_docx.py",
            "--target",
            str(self.target),
            "--facts",
            str(self.facts),
            "--config",
            str(config),
            "--authorized-workspace",
            str(self.workspace),
            "--output-dir",
            str(self.output_dir),
            "--audit-json",
            str(self.audit),
            "--authorization-id",
            "TEST-AUTHORIZATION",
        ]
        def fail_after_partial_write(payload: dict, handle: object, **kwargs: object) -> None:
            handle.write('{"partial":')
            raise OSError("simulated partial audit failure")

        with patch.object(sys, "argv", arguments), patch.object(
            common_module.json, "dump", side_effect=fail_after_partial_write
        ):
            with self.assertRaises(OSError):
                fill_docx_module.main()
        self.assertEqual(list(self.output_dir.glob("*.docx")), [])
        self.assertFalse(self.audit.exists())
        self.assertEqual(list(self.audit_dir.iterdir()), [])

    def test_fill_requires_both_structure_fingerprint_and_file_hash(self) -> None:
        config = self._write_config()
        changed = self.workspace / "same-structure-different-file.docx"
        document = Document(self.target)
        document.core_properties.title = "只改变包元数据"
        document.save(changed)
        self.assertEqual(self._inspection()["template"]["structure_fingerprint"], self._inspection(changed)["template"]["structure_fingerprint"])

        result = run_script(
            "fill_docx.py",
            "--target",
            changed,
            "--facts",
            self.facts,
            "--config",
            config,
            "--authorized-workspace",
            self.workspace,
            "--output-dir",
            self.output_dir,
            "--audit-json",
            self.audit,
            "--authorization-id",
            "TEST-AUTHORIZATION",
            expect_success=False,
        )
        self.assertIn("SHA-256", result.stderr)
        self.assertFalse(self.output_dir.exists())

    def test_fill_rejects_structure_fingerprint_mismatch(self) -> None:
        config = self._write_config()
        changed = self.workspace / "changed-structure.docx"
        document = Document(self.target)
        document.tables[0].add_row()
        document.save(changed)
        result = run_script(
            "fill_docx.py",
            "--target",
            changed,
            "--facts",
            self.facts,
            "--config",
            config,
            "--authorized-workspace",
            self.workspace,
            "--output-dir",
            self.output_dir,
            "--audit-json",
            self.audit,
            "--authorization-id",
            "TEST-AUTHORIZATION",
            expect_success=False,
        )
        self.assertIn("结构指纹", result.stderr)
        self.assertFalse(self.output_dir.exists())

    def test_invalid_mapping_cannot_be_activated(self) -> None:
        draft = self._write_config(status="draft", invalid_mapping=True)
        result = run_script(
            "activate_mapping.py",
            "--draft",
            draft,
            "--authorized-workspace",
            self.workspace,
            "--config-library",
            self.config_library,
            "--approved-by",
            "CRA",
            "--approval-id",
            "TEST-APPROVAL",
            expect_success=False,
        )
        self.assertIn("模板配置不符合 Schema", result.stderr)
        self.assertFalse((self.config_library / "enabled" / draft.name).exists())

    def test_checklist_is_audited_and_tampering_is_rejected(self) -> None:
        config, output_docx, checklist = self._fill_and_build()
        audit = json.loads(self.audit.read_text(encoding="utf-8"))
        self.assertEqual(audit["output_checklist"], str(checklist.resolve()))
        self.assertRegex(audit["output_checklist_sha256"], r"^[0-9a-f]{64}$")

        run_script(
            "validate_outputs.py",
            "--target-input",
            self.target,
            "--output-docx",
            output_docx,
            "--checklist",
            checklist,
            "--config",
            config,
            "--audit-json",
            self.audit,
            "--authorized-workspace",
            self.workspace,
        )

        workbook = load_workbook(checklist)
        workbook["字段核对"]["G2"] = "伪造来源"
        workbook.save(checklist)
        result = run_script(
            "validate_outputs.py",
            "--target-input",
            self.target,
            "--output-docx",
            output_docx,
            "--checklist",
            checklist,
            "--config",
            config,
            "--audit-json",
            self.audit,
            "--authorized-workspace",
            self.workspace,
            expect_success=False,
        )
        self.assertTrue("核对清单" in result.stderr and ("哈希" in result.stderr or "内容" in result.stderr))

    def test_validation_binds_word_path_and_recalculates_audit_from_facts(self) -> None:
        config, output_docx, checklist = self._fill_and_build()
        duplicate = self.workspace / "duplicate.docx"
        shutil.copy2(output_docx, duplicate)
        wrong_path = run_script(
            "validate_outputs.py",
            "--target-input",
            self.target,
            "--output-docx",
            duplicate,
            "--checklist",
            checklist,
            "--config",
            config,
            "--audit-json",
            self.audit,
            "--authorized-workspace",
            self.workspace,
            expect_success=False,
        )
        self.assertIn("输出 Word 路径", wrong_path.stderr)

        audit = json.loads(self.audit.read_text(encoding="utf-8"))
        audit["fields"][0]["source_value"] = "伪造事实"
        self.audit.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
        workbook = load_workbook(checklist)
        workbook["字段核对"]["D2"] = "伪造事实"
        workbook.save(checklist)
        audit["output_checklist_sha256"] = self._sha256(checklist)
        self.audit.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
        tampered = run_script(
            "validate_outputs.py",
            "--target-input",
            self.target,
            "--output-docx",
            output_docx,
            "--checklist",
            checklist,
            "--config",
            config,
            "--audit-json",
            self.audit,
            "--authorized-workspace",
            self.workspace,
            expect_success=False,
        )
        self.assertIn("与输入事实不一致", tampered.stderr)

    def test_validation_rejects_inconsistent_generic_word_output_audit(self) -> None:
        config, output_docx, checklist = self._fill_and_build()
        audit = json.loads(self.audit.read_text(encoding="utf-8"))
        audit["output_form"] = str((self.workspace / "other.docx").resolve())
        self.audit.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
        result = run_script(
            "validate_outputs.py",
            "--target-input",
            self.target,
            "--output-docx",
            output_docx,
            "--checklist",
            checklist,
            "--config",
            config,
            "--audit-json",
            self.audit,
            "--authorized-workspace",
            self.workspace,
            expect_success=False,
        )
        self.assertIn("通用表单输出记录", result.stderr)

    def test_existing_word_sample_preserves_six_statuses_and_never_overwrites(self) -> None:
        sample_target = ROOT / "fixtures" / "form-fill-pilot" / "template" / "XX医院 伦理审查申请表.docx"
        sample_facts = ROOT / "fixtures" / "form-fill-pilot" / "source" / "project-info.xlsx"
        sample_config = (
            ROOT
            / "template-configs"
            / "cra-fill-structured-form"
            / "enabled"
            / "xx-hospital-ethics-review-application-v2.json"
        )
        copied_target = self.workspace / sample_target.name
        copied_facts = self.workspace / sample_facts.name
        copied_config = self.workspace / sample_config.name
        shutil.copy2(sample_target, copied_target)
        shutil.copy2(sample_facts, copied_facts)
        shutil.copy2(sample_config, copied_config)

        outputs: list[Path] = []
        checklists: list[Path] = []
        for run_number in (1, 2):
            audit = self.audit_dir / f"sample-audit-{run_number}.json"
            fill = run_script(
                "fill_docx.py",
                "--target",
                copied_target,
                "--facts",
                copied_facts,
                "--config",
                copied_config,
                "--authorized-workspace",
                self.workspace,
                "--output-dir",
                self.output_dir,
                "--audit-json",
                audit,
                "--authorization-id",
                "TEST-SAMPLE-AUTHORIZATION",
            )
            fill_result = json.loads(fill.stdout)
            self.assertEqual(
                fill_result["status_counts"],
                {"人工保留": 3, "冲突": 0, "已填写": 15, "已核对": 0, "待确认": 1, "缺失": 3},
            )
            outputs.append(Path(fill_result["output_docx"]))
            checklist_result = run_script(
                "build_checklist.py",
                "--audit-json",
                audit,
                "--authorized-workspace",
                self.workspace,
                "--output-dir",
                self.output_dir,
            )
            checklists.append(Path(json.loads(checklist_result.stdout)["checklist"]))
            run_script(
                "validate_outputs.py",
                "--target-input",
                copied_target,
                "--output-docx",
                outputs[-1],
                "--checklist",
                checklists[-1],
                "--config",
                copied_config,
                "--audit-json",
                audit,
                "--authorized-workspace",
                self.workspace,
            )

        self.assertNotEqual(outputs[0], outputs[1])
        self.assertNotEqual(checklists[0], checklists[1])
        self.assertTrue(all(path.exists() for path in [*outputs, *checklists]))

    @staticmethod
    def _sha256(path: Path) -> str:
        import hashlib

        return hashlib.sha256(path.read_bytes()).hexdigest()

    def test_existing_values_are_classified_as_verified_or_conflict_without_overwrite(self) -> None:
        document = Document(self.target)
        document.tables[0].cell(0, 0).text = "原值"
        document.save(self.target)
        config = self._write_config()

        for fact_value, expected_status in (("原值", "已核对"), ("不同值", "冲突")):
            workbook = Workbook()
            sheet = workbook.active
            sheet.append(["字段", "内容"])
            sheet.append(["项目名称", fact_value])
            workbook.save(self.facts)
            audit = self.audit_dir / f"{expected_status}.json"
            result = run_script(
                "fill_docx.py",
                "--target",
                self.target,
                "--facts",
                self.facts,
                "--config",
                config,
                "--authorized-workspace",
                self.workspace,
                "--output-dir",
                self.output_dir,
                "--audit-json",
                audit,
                "--authorization-id",
                "TEST-AUTHORIZATION",
            )
            output = Path(json.loads(result.stdout)["output_docx"])
            self.assertEqual(json.loads(audit.read_text(encoding="utf-8"))["fields"][0]["status"], expected_status)
            self.assertEqual(Document(output).tables[0].cell(0, 0).text, "原值")

    def test_enabled_config_requires_complete_approval(self) -> None:
        config = self._write_config()
        payload = json.loads(config.read_text(encoding="utf-8"))
        payload["approval"]["approval_id"] = None
        config.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        result = run_script(
            "fill_docx.py",
            "--target",
            self.target,
            "--facts",
            self.facts,
            "--config",
            config,
            "--authorized-workspace",
            self.workspace,
            "--output-dir",
            self.output_dir,
            "--audit-json",
            self.audit,
            "--authorization-id",
            "TEST-AUTHORIZATION",
            expect_success=False,
        )
        self.assertIn("审批记录", result.stderr)
        self.assertFalse(self.output_dir.exists())

    def test_fill_requires_nonempty_authorization_and_isolated_output_directory(self) -> None:
        config = self._write_config()
        for authorization_id, output_dir, expected_message in (
            ("", self.output_dir, "授权确认 ID"),
            ("TEST-AUTHORIZATION", self.workspace, "输出目录必须与输入目录隔离"),
        ):
            audit = self.audit_dir / f"audit-{len(expected_message)}.json"
            result = run_script(
                "fill_docx.py",
                "--target",
                self.target,
                "--facts",
                self.facts,
                "--config",
                config,
                "--authorized-workspace",
                self.workspace,
                "--output-dir",
                output_dir,
                "--audit-json",
                audit,
                "--authorization-id",
                authorization_id,
                expect_success=False,
            )
            self.assertIn(expected_message, result.stderr)
            self.assertFalse(audit.exists())

        isolated_output = self.workspace / "isolated-output"
        audit_in_input_dir = self.workspace / "audit-in-input-dir.json"
        result = run_script(
            "fill_docx.py",
            "--target",
            self.target,
            "--facts",
            self.facts,
            "--config",
            config,
            "--authorized-workspace",
            self.workspace,
            "--output-dir",
            isolated_output,
            "--audit-json",
            audit_in_input_dir,
            "--authorization-id",
            "TEST-AUTHORIZATION",
            expect_success=False,
        )
        self.assertIn("输出目录必须与输入目录隔离", result.stderr)
        self.assertFalse(audit_in_input_dir.exists())

    def test_new_template_draft_uses_stable_config_id_and_next_version(self) -> None:
        stable_id = "xx-hospital-ethics-review-application"
        for directory_name, version, status in (("disabled", 1, "disabled"), ("enabled", 2, "enabled")):
            directory = self.config_library / directory_name
            directory.mkdir(parents=True, exist_ok=True)
            (directory / f"{stable_id}-v{version}.json").write_text(
                json.dumps({"config_id": stable_id, "version": version, "status": status}, ensure_ascii=False),
                encoding="utf-8",
            )

        sample_target = ROOT / "fixtures" / "form-fill-pilot" / "template" / "XX医院 伦理审查申请表.docx"
        sample_facts = ROOT / "fixtures" / "form-fill-pilot" / "source" / "project-info.xlsx"
        copied_target = self.workspace / sample_target.name
        copied_facts = self.workspace / sample_facts.name
        shutil.copy2(sample_target, copied_target)
        shutil.copy2(sample_facts, copied_facts)
        result = run_script(
            "draft_mapping.py",
            "--target",
            copied_target,
            "--facts",
            copied_facts,
            "--authorized-workspace",
            self.workspace,
            "--config-library",
            self.config_library,
        )
        draft_path = Path(json.loads(result.stdout)["draft"])
        draft = json.loads(draft_path.read_text(encoding="utf-8"))
        self.assertEqual(draft["config_id"], stable_id)
        self.assertEqual(draft["version"], 3)
        self.assertEqual(draft_path.name, f"{stable_id}-v3.json")


if __name__ == "__main__":
    unittest.main()
